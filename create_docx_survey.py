import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor

def create_survey_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('Survey Đánh Giá Level Dùng AI Tool — Project Manager', level=1)
    
    subtitle = doc.add_paragraph('Mục đích: Xác định PM đang ở level nào về mindset dùng AI và dùng AI tool để quản trị dự án tốt hơn.')
    
    # Instructions
    instructions = doc.add_paragraph('Cách điền: Với mỗi câu, chọn 1 đáp án mô tả đúng nhất tình trạng hiện tại của bạn.')
    
    # --- THÔNG TIN CÁ NHÂN ---
    doc.add_heading('THÔNG TIN CÁ NHÂN & BASELINE', level=2)
    
    personal_questions = [
        "1. Họ và tên",
        "2. Dự án hiện tại",
        "3. Ngày điền (DD/MM/YYYY)",
        "4. Thời gian trung bình để viết 1 PRD / tài liệu yêu cầu (ví dụ: 3 giờ)",
        "5. Thời gian trung bình để viết 1 status report cho stakeholder (ví dụ: 45 phút)",
        "6. Ước tính % công việc PM mỗi tuần hiện có AI hỗ trợ (ví dụ: 30%)"
    ]
    
    for q in personal_questions:
        doc.add_paragraph(q)
        doc.add_paragraph() # Empty paragraph for text response space
        
    # --- PHẦN A -> F: CÂU HỎI TRẮC NGHIỆM ---
    sections = [
        {
            "title": "PHẦN A — Nền tảng dùng AI",
            "questions": [
                {
                    "q": "7. Làm quen công cụ AI: Bạn đã dùng công cụ AI (ChatGPT, Claude, Copilot...) trong công việc như thế nào?",
                    "opts": [
                        "A. Chưa dùng, hoặc chỉ tò mò thử vài lần không có mục đích công việc cụ thể.",
                        "B. Dùng được với prompt/mẫu có sẵn từ thư viện team. (Ví dụ: lấy prompt \"tóm tắt meeting\" có sẵn, dán transcript vào, sửa lại một chút).",
                        "C. Tự viết prompt cho nhu cầu riêng mà không cần mẫu có sẵn. (Ví dụ: tự gõ \"Đóng vai BA, viết 5 user story từ đoạn brief sau, format theo Given-When-Then\" mà không cần ai đưa mẫu).",
                        "D. Dùng AI như một bước mặc định trong mọi việc liên quan. (Ví dụ: mỗi sáng mở AI để review dashboard token, mỗi cuối sprint tự động dùng AI phân tích retro — thành thói quen)."
                    ]
                },
                {
                    "q": "8. Cấu trúc viết prompt: Khi cần AI làm một việc phức tạp, bạn viết prompt như thế nào?",
                    "opts": [
                        "A. Gõ một câu ngắn, không có ngữ cảnh, nếu sai thì thử lại câu khác. (Ví dụ: \"viết PRD giúp tôi\" rồi chờ xem AI ra gì).",
                        "B. Biết thêm role + yêu cầu cơ bản nhưng chưa có cấu trúc rõ ràng. (Ví dụ: \"Bạn là PM, hãy viết PRD cho tính năng đăng nhập.\").",
                        "C. Viết đủ cấu trúc Role + Context + Yêu cầu + Format + Ví dụ, và biết refine lại khi output sai.",
                        "D. Dùng kỹ thuật nâng cao: chain-of-thought (yêu cầu AI suy luận từng bước) hoặc few-shot (đưa ví dụ mẫu để AI bắt chước format)."
                    ]
                }
            ]
        },
        {
            "title": "PHẦN B — Dùng AI cho tác vụ PM",
            "questions": [
                {
                    "q": "9. Tạo PRD bằng AI: Bạn đã dùng AI để soạn PRD (Product Requirement Document) chưa?",
                    "opts": [
                        "A. Chưa từng, vẫn viết tay 100%.",
                        "B. Có thử nhưng phải sửa lại gần hết, mất nhiều thời gian hơn viết tay.",
                        "C. Ra được bản PRD draft đầy đủ cấu trúc từ 1 đoạn brief, trong dưới 30 phút, chỉ cần review và chỉnh sửa nhỏ.",
                        "D. Có checklist review chuẩn trước khi giao team, và đã dùng cách này cho nhiều dự án khác nhau."
                    ]
                },
                {
                    "q": "10. Tạo User Story & Acceptance Criteria bằng AI: Bạn dùng AI để viết User Story và AC như thế nào?",
                    "opts": [
                        "A. Chưa dùng AI cho việc này.",
                        "B. Có dùng nhưng chưa biết cách review — copy nguyên output vào Jira mà không kiểm tra kỹ.",
                        "C. Dùng AI generate User Story từ PRD, generate AC từ Story, và tự phát hiện được chỗ AI viết mơ hồ hoặc thiếu edge case trước khi giao team.",
                        "D. Có quy trình chuẩn: brief → AI draft → tự review theo checklist riêng → AI refine lần 2 → giao team."
                    ]
                },
                {
                    "q": "11. Status Report cho Stakeholder: Bạn dùng AI để viết báo cáo tiến độ cho stakeholder như thế nào?",
                    "opts": [
                        "A. Chưa dùng, tự viết tay hoàn toàn.",
                        "B. Dùng AI viết nháp nhưng vẫn còn thuật ngữ kỹ thuật, stakeholder non-technical đọc không hiểu, phải hỏi lại.",
                        "C. Dùng AI translate số liệu kỹ thuật sang business value dễ hiểu, report 1 trang, stakeholder hiểu ngay không cần hỏi thêm.",
                        "D. Ngoài report viết, còn dùng AI hỗ trợ chuẩn bị talking points cho các câu hỏi khó từ khách hàng, và có thể present bằng video/slide."
                    ]
                },
                {
                    "q": "12. Theo dõi Token Cost bằng AI: Bạn dùng AI để theo dõi token cost như thế nào?",
                    "opts": [
                        "A. Không theo dõi token cost, hoặc để việc này cho Dev/Tech Lead.",
                        "B. Có đọc dashboard nhưng không phân tích, chỉ nhìn con số tổng.",
                        "C. Dùng AI tóm tắt dashboard, nhận diện được task/agent nào đang dùng nhiều token hơn dự tính, phân biệt được spike bình thường và spike đáng lo.",
                        "D. Dùng AI để so sánh token/feature qua nhiều sprint, phát hiện pattern, và tự forecast budget 1-2 sprint tới từ baseline."
                    ]
                },
                {
                    "q": "13. Prompt Library cá nhân: Bạn có lưu lại prompt để dùng lại không?",
                    "opts": [
                        "A. Không lưu, mỗi lần cần lại gõ lại từ đầu.",
                        "B. Có lưu vài prompt rời rạc (note, chat history) nhưng không có tổ chức, khó tìm lại.",
                        "C. Có thư viện prompt cá nhân ≥ 10 prompt, có ghi chú (task nào, dùng tool gì, khi nào dùng), tìm lại dễ dàng.",
                        "D. Ngoài thư viện cá nhân, đã chia sẻ/đóng góp prompt cho team dùng chung, có ít nhất 1 template được team khác áp dụng thực tế."
                    ]
                },
                {
                    "q": "14. AI trong Workflow Sprint hàng ngày: AI có mặt trong công việc sprint của bạn ở mức nào?",
                    "opts": [
                        "A. Chỉ dùng AI khi \"chợt nhớ ra\", không phải việc thường xuyên.",
                        "B. Dùng AI cho 1-2 việc cố định (vd: luôn dùng AI viết status report) nhưng chưa mở rộng ra các bước khác.",
                        "C. AI có mặt xuyên suốt sprint cycle: hỗ trợ Sprint Planning, Backlog Refinement, Retrospective.",
                        "D. Đã đo được: workflow có AI tiết kiệm ≥ 30% thời gian so với làm tay (có số liệu cụ thể, không phải ước chừng)."
                    ]
                },
                {
                    "q": "15. Phân tích dữ liệu & rút insight bằng AI: Bạn dùng AI để phân tích dữ liệu (feedback, metrics, sprint data) như thế nào?",
                    "opts": [
                        "A. Chưa dùng AI cho việc phân tích, chỉ đọc số liệu thô.",
                        "B. Dùng AI tóm tắt dữ liệu nhưng dừng ở việc liệt kê lại, chưa có insight hay đề xuất hành động.",
                        "C. Dùng AI đi từ raw data ra được insight cụ thể và đề xuất hành động rõ ràng.",
                        "D. Kết hợp phân tích insight với forecast, dùng kết quả để điều chỉnh kế hoạch trước khi vấn đề xảy ra."
                    ]
                },
                {
                    "q": "16. Chọn đúng tool cho đúng việc: Bạn có phân biệt được nên dùng tool AI nào cho việc gì không?",
                    "opts": [
                        "A. Chỉ biết 1 tool, dùng cho mọi việc bất kể phù hợp hay không.",
                        "B. Biết vài tool khác nhau (ChatGPT, Claude, Notion AI, Jira AI...) nhưng chọn theo thói quen, chưa so sánh rõ ưu nhược điểm.",
                        "C. Biết so sánh và chọn đúng tool theo use case cụ thể. (Ví dụ: dùng Jira AI cho backlog; dùng Claude/ChatGPT cho viết PRD).",
                        "D. Từng tham gia đánh giá/thử nghiệm (pilot) tool mới trước khi đề xuất cho team, có tiêu chí đánh giá rõ ràng."
                    ]
                }
            ]
        },
        {
            "title": "PHẦN C — Mindset & Phán đoán khi dùng AI",
            "questions": [
                {
                    "q": "17. Biết khi nào KHÔNG nên dùng AI: Bạn có phân biệt được việc nào nên và không nên giao cho AI không?",
                    "opts": [
                        "A. Nghĩ AI dùng được cho gần như mọi việc, cứ thử là được.",
                        "B. Mơ hồ cảm thấy có việc AI không hợp, nhưng chưa rõ ranh giới, thường vẫn dùng rồi mới biết.",
                        "C. Xác định rõ loại việc không nên giao AI và chủ động không dùng. (Ví dụ: quyết định nhạy cảm về nhân sự/khách hàng, số liệu cần chính xác tuyệt đối, cam kết pháp lý).",
                        "D. Hướng dẫn được người khác phân biệt task nên/không nên dùng AI, và giải thích được lý do."
                    ]
                },
                {
                    "q": "18. Review output & trách nhiệm với hallucination: Bạn xử lý kết quả AI đưa ra như thế nào?",
                    "opts": [
                        "A. Tin tưởng và dùng thẳng output AI, ít khi kiểm tra lại.",
                        "B. Có đọc lại nhưng chủ yếu kiểm tra câu chữ, không đối chiếu số liệu/nguồn.",
                        "C. Luôn review, đối chiếu số liệu và logic với nguồn thật, sửa chỗ AI bịa trước khi dùng; hiểu rõ con người chịu trách nhiệm cuối cùng.",
                        "D. Có checklist/quy trình verify chuẩn cho output AI, và hướng dẫn team cùng áp dụng để tránh hallucination."
                    ]
                }
            ]
        },
        {
            "title": "PHẦN D — Bảo mật & Tuân thủ",
            "questions": [
                {
                    "q": "19. Bảo mật dữ liệu khi dùng AI: Bạn kiểm soát dữ liệu đưa vào công cụ AI như thế nào?",
                    "opts": [
                        "A. Chưa nghĩ tới rủi ro này — đưa mọi thứ vào AI kể cả spec, dữ liệu khách hàng, source code nhạy cảm.",
                        "B. Biết là có rủi ro bảo mật nhưng chưa rõ ranh giới dữ liệu nào được/không được đưa vào.",
                        "C. Phân loại được và tuân thủ: biết dữ liệu nào KHÔNG được đưa vào tool AI công cộng, chỉ dùng tool/kênh được công ty cho phép cho dữ liệu nhạy cảm.",
                        "D. Tham gia xây dựng hoặc phổ biến policy về acceptable use / phân loại dữ liệu cho team."
                    ]
                }
            ]
        },
        {
            "title": "PHẦN E — Kết nối, Skill, Tự động hóa",
            "questions": [
                {
                    "q": "20. Kết nối AI với dữ liệu/công cụ khác (MCP): Bạn có dùng AI kết nối trực tiếp với công cụ khác (Jira, Backlog, Slack, Notion...) qua MCP/connector không?",
                    "opts": [
                        "A. Chưa biết MCP/connector là gì — luôn copy-paste dữ liệu thủ công vào AI.",
                        "B. Biết khái niệm MCP/connector, nghe đồng nghiệp hoặc admin nhắc tới, nhưng chưa tự dùng.",
                        "C. Đã tự dùng AI kết nối với ít nhất 1 công cụ qua MCP để lấy/cập nhật dữ liệu trực tiếp.",
                        "D. Kết hợp nhiều MCP/connector trong cùng một quy trình để tự động hóa việc lặp lại."
                    ]
                },
                {
                    "q": "21. Dùng/tạo Skill: Bạn đã dùng hoặc tự tạo \"Skill\" (quy trình AI đóng gói sẵn) chưa?",
                    "opts": [
                        "A. Chưa biết khái niệm Skill — mỗi lần làm việc lặp lại vẫn gõ prompt lại từ đầu.",
                        "B. Có dùng Skill có sẵn (do công ty/đồng nghiệp tạo) cho một vài việc quen thuộc, nhưng chưa tự tạo Skill riêng.",
                        "C. Tự tạo được ít nhất 1 Skill cho việc lặp lại của riêng mình.",
                        "D. Skill do bạn tạo đang được người khác trong team dùng lại thật."
                    ]
                },
                {
                    "q": "22. Bộ công cụ SDLC nội bộ (Spec → RAG → BD/DD/SRS → Testcase → Code): Bạn đã dùng bộ công cụ này ở mức nào?",
                    "opts": [
                        "A. Chưa từng nghe nói về bộ công cụ này.",
                        "B. Có nghe/được giới thiệu, nhưng chưa tự dùng vào dự án thực tế nào.",
                        "C. Đã dùng thử trong ít nhất 1 dự án — đưa spec vào storage, để hệ thống RAG phân tích và sinh ra BD/DD/SRS, testcase.",
                        "D. Đã dùng thường xuyên trong dự án đang chạy, tự tin chuẩn hóa input để RAG phân tích ra output chất lượng cao."
                    ]
                },
                {
                    "q": "23. Custom tool / Automation: Bạn đã từng xây dựng công cụ AI riêng hoặc automation chưa?",
                    "opts": [
                        "A. Chưa từng, chỉ dùng tool có sẵn qua giao diện chat thông thường.",
                        "B. Có nghe/tìm hiểu về Custom GPT, Zapier, Make... nhưng chưa tự làm.",
                        "C. Đã tự xây 1 Custom GPT/Agent hoặc automation đơn giản và đang dùng thật cho công việc.",
                        "D. Công cụ/automation do bạn xây đang được cả team dùng thực tế."
                    ]
                }
            ]
        },
        {
            "title": "PHẦN F — Lan tỏa & Chiến lược",
            "questions": [
                {
                    "q": "24. Xây dựng & chia sẻ template cho team: Bạn đã đóng góp gì cho việc chuẩn hóa AI trong team chưa?",
                    "opts": [
                        "A. Chưa từng chia sẻ, chỉ dùng cho việc cá nhân.",
                        "B. Có chia sẻ 1-2 prompt cho đồng nghiệp khi họ hỏi, nhưng không chủ động.",
                        "C. Chủ động xây dựng prompt template chuẩn kèm hướng dẫn sử dụng, có ít nhất 1 template team đang dùng thật.",
                        "D. Đã hướng dẫn/mentor được ít nhất 1 PM khác áp dụng AI vào công việc."
                    ]
                },
                {
                    "q": "25. Chiến lược AI Tool cấp tổ chức: Bạn có tham gia vào việc định hướng dùng AI tool ở quy mô lớn hơn nhóm mình không?",
                    "opts": [
                        "A. Chưa, chỉ quan tâm ở phạm vi công việc cá nhân/dự án mình.",
                        "B. Có đóng góp ý kiến khi được hỏi nhưng chưa chủ động đề xuất.",
                        "C. Đã chủ động đề xuất roadmap/ngân sách adoption AI tool cho nhóm/phòng ban.",
                        "D. Đã xây dựng chiến lược AI tool được lãnh đạo phê duyệt và đang triển khai ở cấp SBU/công ty."
                    ]
                }
            ]
        }
    ]
    
    for sec in sections:
        doc.add_heading(sec["title"], level=2)
        for item in sec["questions"]:
            doc.add_paragraph(item["q"])
            for opt in item["opts"]:
                doc.add_paragraph(opt)
            doc.add_paragraph() # Blank line between questions
            
    # --- PHẦN G: CÂU HỎI MỜ ---
    doc.add_heading('PHẦN G — Câu hỏi mở (Điền tự do)', level=2)
    
    open_qs = [
        "26. Khó khăn lớn nhất của bạn khi đưa AI vào công việc PM là gì?",
        "27. Kỹ năng hoặc công cụ AI nào bạn muốn học nhất trong 3 tháng tới?",
        "28. Bạn đã ứng dụng thành công công cụ AI nào vào công việc PM chưa? Bạn có sẵn sàng chia sẻ/mentor lại cho team không?",
        "29. Với bộ công cụ SDLC nội bộ (spec → RAG → BD/DD/SRS → testcase → code), bạn thấy vướng mắc hoặc rủi ro đặc thù gì?"
    ]
    for oq in open_qs:
        doc.add_paragraph(oq)
        doc.add_paragraph()
        
    doc.add_paragraph("30. Bạn muốn học theo hình thức nào? (có thể chọn nhiều)")
    multi_opts = [
        "A. Workshop offline",
        "B. Online",
        "C. Tự học + Q&A",
        "D. OJT kèm mentor",
        "E. Linh hoạt theo lịch dự án"
    ]
    for mo in multi_opts:
        doc.add_paragraph(mo)
        
    output_path = os.path.join(os.getcwd(), 'survey-ms-forms-import.docx')
    doc.save(output_path)
    print(f"File created successfully: {output_path}")

if __name__ == "__main__":
    create_survey_docx()
